# --- Final Production Code ---

import io
import os
import requests
import concurrent.futures
import logging
import time
from typing import List, Dict, Any, Optional

# Qdrant
import qdrant_client
from qdrant_client.http.models import Distance, VectorParams, PointStruct, CollectionStatus, OptimizersConfigDiff, HnswConfigDiff

# LangChain components
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Qdrant
from langchain_openai import OpenAIEmbeddings # Use OpenAI embeddings

# File Parsers
import pdfplumber
from docx import Document as DocxDocument

# Setup robust logging for production
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("kb_indexer.log"), # Log to a file
        logging.StreamHandler() # Log to console
    ]
)
logger = logging.getLogger(__name__)

# --- Constants ---
# Dimension for OpenAI's text-embedding-3-small
OPENAI_EMBEDDING_DIMENSION = 1536
# Default Qdrant distance metric suitable for OpenAI embeddings
QDRANT_DISTANCE_METRIC = Distance.COSINE

# --- Helper Function to Parse Files ---
def parse_file_content(content: bytes, filename: str) -> Optional[str]:
    """Parses content based on filename extension."""
    logger.debug(f"Parsing file: {filename}")
    try:
        file_ext = os.path.splitext(filename)[1].lower()
        if file_ext == '.pdf':
            text = ""
            # Using BytesIO for in-memory processing
            with io.BytesIO(content) as pdf_stream:
                with pdfplumber.open(pdf_stream) as pdf:
                    # Handle potential empty pages or extract errors gracefully
                    for page_num, page in enumerate(pdf.pages):
                         page_text = page.extract_text()
                         if page_text:
                             text += page_text + "\n"
                         else:
                             logger.debug(f"No text extracted from page {page_num + 1} of {filename}")
            logger.debug(f"Finished parsing PDF: {filename}")
            return text.strip() if text else None # Return None if empty
        elif file_ext == '.docx':
             # Using BytesIO for in-memory processing
            with io.BytesIO(content) as docx_stream:
                doc = DocxDocument(docx_stream)
                text = "\n".join(para.text for para in doc.paragraphs if para.text)
            logger.debug(f"Finished parsing DOCX: {filename}")
            return text.strip() if text else None # Return None if empty
        elif file_ext == '.txt':
            # Attempt decoding with UTF-8, replace errors
            text = content.decode('utf-8', errors='replace')
            logger.debug(f"Finished parsing TXT: {filename}")
            return text.strip() if text else None # Return None if empty
        else:
            logger.warning(f"Unsupported file type: {filename}. Skipping.")
            return None
    except Exception as e:
        logger.error(f"Error parsing file {filename}: {e}", exc_info=True) # Log traceback
        return None

# --- Helper Function to Process a Single File ---
def process_single_file(file_url: str, splitter: RecursiveCharacterTextSplitter) -> List[Document]:
    """Downloads, parses, and chunks a single file."""
    documents = []
    try:
        logger.info(f"Processing file: {file_url}")
        # Increased timeout for potentially large files
        response = requests.get(file_url, timeout=120)
        response.raise_for_status()  # Raise HTTPError for bad responses (4xx or 5xx)

        # More robust filename extraction from URL
        try:
            # Attempt to get filename from Content-Disposition header first
            content_disposition = response.headers.get('Content-Disposition')
            if content_disposition:
                import re
                fname = re.findall('filename="?(.+)"?', content_disposition)
                if fname:
                    filename = fname[0]
                else: # Fallback if header is malformed
                    filename = os.path.basename(file_url.split('?')[0].split('/')[-1])

            else: # Fallback to URL path component
                 filename = os.path.basename(file_url.split('?')[0].split('/')[-1])

            # Handle cases where filename might be empty or just '/'
            if not filename or filename == '/':
                filename = f"unknown_file_{time.time()}" # Generate a placeholder
                logger.warning(f"Could not determine filename for {file_url}, using placeholder: {filename}")

        except Exception as e:
             logger.warning(f"Error extracting filename for {file_url}, using fallback. Error: {e}")
             filename = f"unknown_file_{time.time()}" # Generate a placeholder


        content = response.content
        if not content:
             logger.warning(f"Downloaded file {filename} ({file_url}) is empty. Skipping.")
             return []

        extracted_text = parse_file_content(content, filename)

        if extracted_text:
            # Use the splitter to create chunks
            chunks = splitter.split_text(extracted_text)
            # Create LangChain Document objects with metadata
            for i, chunk in enumerate(chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "source": file_url,
                        "filename": filename,
                        "chunk_index": i,
                        "indexed_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()) # Add timestamp
                    }
                )
                documents.append(doc)
            logger.info(f"Successfully chunked {filename} into {len(documents)} documents.")
        else:
             logger.warning(f"No text could be extracted from {filename} ({file_url}).")


    except requests.exceptions.Timeout:
        logger.error(f"Timeout while downloading {file_url}")
    except requests.exceptions.RequestException as e:
        logger.error(f"Failed to download or access {file_url}: {e}")
    except Exception as e:
        # Catch any other unexpected errors during processing
        logger.error(f"An unexpected error occurred processing {file_url}: {e}", exc_info=True)

    return documents

# --- Main Indexer Function ---
def KB_indexer(
    file_urls: List[str],
    qdrant_url: str,
    qdrant_api_key: Optional[str], # API key is optional
    collection_name: str,
    openai_api_key: Optional[str] = None, # Allow passing key directly if needed, but env var is preferred
    force_recreate_collection: bool = False, # Option to recreate collection
    max_workers: int = 10 # Max parallel download/parse threads
) -> bool:
    """
    Indexes files from URLs into a Qdrant collection using OpenAI embeddings
    (text-embedding-3-small) and parallel processing.

    Relies on the OPENAI_API_KEY environment variable if openai_api_key argument is None.

    Args:
        file_urls: List of URLs pointing to the files (PDF, DOCX, TXT).
        qdrant_url: URL of the Qdrant instance (e.g., "http://localhost:6333").
        qdrant_api_key: API key for Qdrant Cloud or secured instances. None if not needed.
        collection_name: Name of the Qdrant collection (e.g., user_email+gpt_name).
                         Should be unique per knowledge base.
        openai_api_key: Optional OpenAI API Key. If None, reads from OPENAI_API_KEY env var.
        force_recreate_collection: If True, deletes the collection if it exists before indexing.
                                   Use with caution in production.
        max_workers: Maximum number of parallel threads for downloading/parsing files.

    Returns:
        True if the indexing process completed successfully (even if some files failed).
        False if a critical error occurred (Qdrant connection, collection setup, embedding failure).
    """
    start_time = time.time()
    logger.info(f"Starting indexing process for collection: {collection_name}")
    logger.info(f"Processing {len(file_urls)} file(s).")
    if not file_urls:
        logger.warning("No file URLs provided. Exiting.")
        return True # Nothing to do, considered successful

    # 1. Initialize OpenAI Embeddings
    try:
        # Reads OPENAI_API_KEY env var by default if api_key is None
        # Specify the model explicitly
        embedding_model = OpenAIEmbeddings(
            model="text-embedding-3-small",
            api_key=openai_api_key, # Pass key if provided, else None (reads from env)
            # Consider adding request_timeout if needed
            # request_timeout=60,
            chunk_size=1000 # Default chunk size for embeddings API requests
        )
        # Optional: Perform a quick test embed to catch auth errors early
        # embedding_model.embed_query("test")
        logger.info("OpenAI embedding model initialized for 'text-embedding-3-small'.")
    except Exception as e:
        logger.error(f"Failed to initialize OpenAI embedding model: {e}", exc_info=True)
        logger.error("Ensure OPENAI_API_KEY environment variable is set correctly or passed as an argument.")
        return False

    # 2. Initialize Qdrant Client
    try:
        client = qdrant_client.QdrantClient(
            url=qdrant_url,
            api_key=qdrant_api_key,
            timeout=60.0 # Increase timeout for potentially long operations
            # prefer_grpc=True, # Consider enabling for performance if network allows and server supports it
        )
        client.health_check() # Verify connection
        logger.info(f"Qdrant client initialized and connection verified to {qdrant_url}.")
    except Exception as e:
        logger.error(f"Failed to initialize Qdrant client at {qdrant_url}: {e}", exc_info=True)
        return False # Indicate failure

    # 3. Check/Create Qdrant Collection
    try:
        vector_size = OPENAI_EMBEDDING_DIMENSION # Use constant
        collection_exists = False
        try:
            collection_info = client.get_collection(collection_name=collection_name)
            if collection_info:
                collection_exists = True
                logger.info(f"Collection '{collection_name}' already exists.")
                # Validate existing collection parameters
                existing_vector_size = -1
                if isinstance(collection_info.vectors_config, qdrant_client.http.models.VectorsConfig):
                    existing_vector_size = collection_info.vectors_config.params.size
                elif isinstance(collection_info.vectors_config, dict): # Handle named vectors potentially
                    # Assuming default vector name if using Langchain wrapper default
                    default_vector_name = 'content' # This assumption might need adjustment
                    if default_vector_name in collection_info.vectors_config:
                         existing_vector_size = collection_info.vectors_config[default_vector_name].params.size


                if force_recreate_collection:
                    logger.warning(f"Recreating collection '{collection_name}' as force_recreate_collection is True.")
                    client.delete_collection(collection_name=collection_name, timeout=120)
                    # Short pause after deletion before recreation
                    time.sleep(2)
                    collection_exists = False
                elif existing_vector_size != vector_size:
                     logger.error(f"Collection '{collection_name}' exists but has incorrect vector size "
                                  f"({existing_vector_size})! Expected {vector_size} for 'text-embedding-3-small'. "
                                  f"Cannot proceed. Use force_recreate_collection=True or delete/migrate manually.")
                     return False # Stop for safety

        except qdrant_client.http.exceptions.UnexpectedResponse as e:
             # Specifically catch 404 Not Found
             if hasattr(e, 'status_code') and e.status_code == 404:
                 logger.info(f"Collection '{collection_name}' does not exist. Will be created.")
                 collection_exists = False
             else: # Re-raise other unexpected errors
                 raise e
        except Exception as e:
             # Catch potential connection errors or other client issues during get_collection
             logger.error(f"Error checking collection '{collection_name}': {e}", exc_info=True)
             return False


        if not collection_exists:
            logger.info(f"Creating collection '{collection_name}' with vector size {vector_size} and distance {QDRANT_DISTANCE_METRIC}.")
            client.create_collection(
                collection_name=collection_name,
                vectors_config=VectorParams(size=vector_size, distance=QDRANT_DISTANCE_METRIC),
                # Optional: Add optimizer/indexing parameters for production performance
                optimizers_config=OptimizersConfigDiff(memmap_threshold=20000), # Example: Adjust based on expected data size
                hnsw_config=HnswConfigDiff(m=16, ef_construct=100) # Example: HNSW parameters, tune as needed
            )
            # Wait briefly for collection to likely become active
            time.sleep(2)
            logger.info(f"Collection '{collection_name}' created.")

    except Exception as e:
        logger.error(f"Error during Qdrant collection setup for '{collection_name}': {e}", exc_info=True)
        return False

    # 4. Initialize Text Splitter
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=120, # Overlap helps maintain context between chunks
        length_function=len,
        is_separator_regex=False, # Use default separators like "\n\n", "\n", " "
    )

    # 5. Process Files in Parallel
    all_documents: List[Document] = []
    processed_files_count = 0
    failed_files: List[str] = []

    # Adjust max_workers based on system resources and network latency
    actual_max_workers = min(max_workers, len(file_urls) if file_urls else 1)
    logger.info(f"Using up to {actual_max_workers} parallel workers for file processing.")

    with concurrent.futures.ThreadPoolExecutor(max_workers=actual_max_workers) as executor:
        future_to_url = {executor.submit(process_single_file, url, text_splitter): url for url in file_urls}

        for future in concurrent.futures.as_completed(future_to_url):
            url = future_to_url[future]
            try:
                documents_from_file = future.result() # Get the list of Document objects
                if documents_from_file:
                    all_documents.extend(documents_from_file)
                    processed_files_count += 1
                else:
                    # Log files that yielded no documents (might indicate parse failure or empty file)
                    logger.warning(f"File {url} resulted in zero documents.")
                    # We might still count it as processed, but failed to extract content
                    failed_files.append(url)

            except Exception as e:
                # Log exceptions raised from within process_single_file if not caught there
                logger.error(f"Error retrieving result for URL {url}: {e}", exc_info=True)
                failed_files.append(url)


    logger.info(f"Finished processing files. Successful: {processed_files_count}, Failed/Empty: {len(failed_files)}.")
    if failed_files:
        logger.warning(f"URLs that failed or yielded no content: {failed_files}")

    if not all_documents:
        logger.warning("No documents were generated from the provided URLs. Nothing to index.")
        total_time = time.time() - start_time
        logger.info(f"Indexing process completed in {total_time:.2f} seconds (no documents indexed).")
        # Return True because the process itself didn't hit a critical error, just no data
        return True

    logger.info(f"Total document chunks generated from all files: {len(all_documents)}")

    # 6. Embed and Upsert documents into Qdrant
    try:
        logger.info(f"Starting embedding and upserting into collection '{collection_name}'...")

        # Use Langchain's Qdrant wrapper for convenience
        qdrant_vector_store = Qdrant(
            client=client,
            collection_name=collection_name,
            embeddings=embedding_model, # Pass the initialized OpenAIEmbeddings instance
            # Increase batch size for OpenAI, as their API handles multiple texts per request efficiently
            # Adjust based on typical chunk size and OpenAI rate limits / recommendations
            batch_size=128,
            # timeout=120.0 # Consider increasing timeout for large embedding batches
        )

        # add_documents handles embedding calls to OpenAI and upserting to Qdrant
        # It returns the list of Qdrant point IDs added
        ids = qdrant_vector_store.add_documents(all_documents, ids=None) # Let Qdrant generate IDs

        logger.info(f"Successfully upserted {len(ids)} vectors into collection '{collection_name}'.")
        total_time = time.time() - start_time
        logger.info(f"Indexing process completed successfully in {total_time:.2f} seconds.")
        return True # Indicate success

    except Exception as e:
        # Catch potential errors during embedding (e.g., OpenAI API errors) or upserting
        logger.error(f"CRITICAL: Failed to embed and upsert documents into Qdrant collection '{collection_name}': {e}", exc_info=True)
        total_time = time.time() - start_time
        logger.error(f"Indexing process failed after {total_time:.2f} seconds during embedding/upserting.")
        return False # Indicate failure

# --- How to Use (Instructions incorporated in comments and example) ---

if __name__ == '__main__':

    print("--- KB Indexer Example ---")

    # --- 1. Prerequisites ---
    #    - Python 3.8+ installed
    #    - Access to a Qdrant instance (local or cloud)
    #    - An OpenAI API Key

    # --- 2. Installation ---
    #    Run in your terminal:
    #    pip install qdrant-client "langchain>=0.1.0" langchain-community langchain-text-splitters langchain-openai openai tiktoken pdfplumber python-docx requests

    # --- 3. Environment Variables ---
    #    Set your OpenAI API Key. It's recommended to use environment variables for secrets.
    #    Linux/macOS: export OPENAI_API_KEY='your-openai-api-key'
    #    Windows (cmd): set OPENAI_API_KEY=your-openai-api-key
    #    Windows (PowerShell): $env:OPENAI_API_KEY='your-openai-api-key'
    #    Alternatively, you can pass the key directly to the KB_indexer function via the `openai_api_key` argument,
    #    but using environment variables is generally more secure.

    # --- 4. Configuration ---
    # Replace these placeholders with your actual values

    # Qdrant Configuration
    QDRANT_URL = os.environ.get("QDRANT_URL", "http://localhost:6333") # Your Qdrant instance URL
    QDRANT_API_KEY = os.environ.get("QDRANT_API_KEY", None) # Your Qdrant API Key (if using Qdrant Cloud or authentication)

    # File URLs to index (replace with URLs accessible by the script)
    # Ensure these are direct links to PDF, DOCX, or TXT files.
    FILE_URLS = [
        "https://www.w3.org/WAI/ER/tests/xhtml/testfiles/resources/pdf/dummy.pdf", # Example PDF
        "https://www.unm.edu/~unmvclib/powerpoint/pptexamples/sample.docx", # Example DOCX (check if link works)
        "https://raw.githubusercontent.com/google/gemini-api/main/README.md", # Example TXT/MD (parsed as TXT)
        "https://invalid-url-that-will-fail.xyz/document.pdf" # Example of a failing URL
    ]

    # Collection Name Construction (example based on prompt)
    USER_EMAIL = "test@example.com"
    GPT_NAME = "my_openai_rag_gpt"
    # Sanitize email and GPT name to create a valid collection name
    # Qdrant collection names must start with a letter and contain only letters, numbers, and underscores.
    sanitized_email = ''.join(c if c.isalnum() else '_' for c in USER_EMAIL)
    sanitized_gpt_name = ''.join(c if c.isalnum() else '_' for c in GPT_NAME)
    # Ensure it starts with a letter if the sanitized email doesn't
    collection_prefix = "kb" if not sanitized_email or not sanitized_email[0].isalpha() else ""
    dynamic_collection_name = f"{collection_prefix}_{sanitized_email}_{sanitized_gpt_name}"
    # Ensure name is not excessively long if needed
    dynamic_collection_name = dynamic_collection_name[:63] # Example length limit if necessary

    # --- 5. Running the Indexer ---
    print(f"Target Qdrant URL: {QDRANT_URL}")
    print(f"Target Collection Name: {dynamic_collection_name}")
    print(f"Number of files to process: {len(FILE_URLS)}")

    # Check if OpenAI API key is available (optional check, the function handles it)
    openai_key_present = bool(os.environ.get("OPENAI_API_KEY"))
    if not openai_key_present:
         print("\nWARNING: OPENAI_API_KEY environment variable not found. Make sure it's set.")
         # You could add logic here to prompt for the key or read from a config file if needed
    else:
        print("OpenAI API key found in environment.")

    # Execute the main function
    # Set force_recreate_collection=True ONLY if you want to clear the collection first.
    # Set force_recreate_collection=False to append/update documents.
    success = KB_indexer(
        file_urls=FILE_URLS,
        qdrant_url=QDRANT_URL,
        qdrant_api_key=QDRANT_API_KEY,
        collection_name=dynamic_collection_name,
        openai_api_key=None, # Set to your key string if not using env var, e.g., "sk-..."
        force_recreate_collection=True, # Be careful with True in production!
        max_workers=5 # Adjust based on your machine/network
    )

    # --- 6. Result ---
    if success:
        print(f"\n--- KB Indexer finished successfully for collection '{dynamic_collection_name}' ---")
    else:
        print(f"\n--- KB Indexer failed for collection '{dynamic_collection_name}'. Check 'kb_indexer.log' for details. ---")